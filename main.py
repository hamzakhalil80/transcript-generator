# main.py
from fastapi import FastAPI, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
import yt_dlp, requests, os, random, time
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from typing import List, Tuple, Optional

app = FastAPI()

# ✅ Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # adjust if you want to restrict to frontend domain
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory list of preferred cookies (keeps working cookies at the front)
preferred_cookies: List[str] = []


# -------- Helper Function ----------
def _find_cookie_files() -> List[str]:
    """Return cookie files present in cwd that match cookies_*.txt"""
    return sorted([f for f in os.listdir(".") if f.startswith("cookies_") and f.endswith(".txt")])


def _make_ordered_cookie_list() -> List[str]:
    """Return cookie files in preferred order: preferred first, then shuffled rest."""
    cookie_files = _find_cookie_files()
    ordered = []
    # Add existing preferred cookies first (only if they still exist)
    for c in preferred_cookies:
        if c in cookie_files and c not in ordered:
            ordered.append(c)
    # Add remaining cookie files shuffled
    remaining = [c for c in cookie_files if c not in ordered]
    random.shuffle(remaining)
    ordered.extend(remaining)
    return ordered


def _update_preferred(cookie_file: str):
    """Move a working cookie_file to the front of preferred_cookies (no duplicates)."""
    global preferred_cookies
    if cookie_file in preferred_cookies:
        preferred_cookies.remove(cookie_file)
    preferred_cookies.insert(0, cookie_file)
    # Keep list bounded to number of cookie files (avoid growing)
    maxlen = max(6, len(_find_cookie_files()))
    preferred_cookies = preferred_cookies[:maxlen]


def _extract_with_options(url: str, ydl_opts: dict):
    """Helper to call yt_dlp and return info (raises exceptions through)."""
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        return ydl.extract_info(url, download=False)


def extract_transcript(video_id: str) -> Tuple[dict, Optional[List[dict]]]:
    """
    Fetch YouTube transcript (tries cookie rotation).
    Returns (info, transcript_list) where transcript_list is [{'start', 'text'}, ...] or None.
    """
    url = f"https://www.youtube.com/watch?v={video_id}"

    cookie_files = _make_ordered_cookie_list()

    # Try each cookie file first (if any)
    for cookie_file in cookie_files:
        ydl_opts = {
            "skip_download": True,
            "quiet": True,
            "cookiefile": cookie_file,
        }
        try:
            info = _extract_with_options(url, ydl_opts)

            subs = {}
            if info.get("subtitles"):
                subs.update(info["subtitles"])
            if info.get("automatic_captions"):
                subs.update(info["automatic_captions"])

            if not subs:
                # cookie worked but no subtitles — try next cookie
                continue

            en_key = next((k for k in subs.keys() if k.startswith("en")), None)
            if not en_key:
                continue

            transcript_url = subs[en_key][0]["url"]
            res = requests.get(transcript_url, timeout=15)
            res.raise_for_status()
            transcript_json = res.json()
            events = transcript_json.get("events", [])

            transcript = [
                {"start": e["tStartMs"] / 1000, "text": seg["utf8"]}
                for e in events if "segs" in e
                for seg in e["segs"]
                if seg.get("utf8", "").strip()
            ]

            # success: promote this cookie for next requests
            _update_preferred(cookie_file)
            return info, transcript

        except Exception as e:
            # log and continue trying next cookie
            print(f"[WARN] cookie '{cookie_file}' failed: {e}")
            continue

    # Last fallback: try without cookies (public access)
    try:
        info = _extract_with_options(url, {"skip_download": True, "quiet": True})

        subs = {}
        if info.get("subtitles"):
            subs.update(info["subtitles"])
        if info.get("automatic_captions"):
            subs.update(info["automatic_captions"])

        if not subs:
            return info, None

        en_key = next((k for k in subs.keys() if k.startswith("en")), None)
        if not en_key:
            return info, None

        transcript_url = subs[en_key][0]["url"]
        res = requests.get(transcript_url, timeout=15)
        res.raise_for_status()
        transcript_json = res.json()
        events = transcript_json.get("events", [])

        transcript = [
            {"start": e["tStartMs"] / 1000, "text": seg["utf8"]}
            for e in events if "segs" in e
            for seg in e["segs"]
            if seg.get("utf8", "").strip()
        ]

        return info, transcript

    except Exception as e:
        # Bubble up meaningful error
        raise RuntimeError(str(e)) from e


# -------- Merge transcript into natural paragraphs ----------
def merge_paragraphs(transcript, max_chars=300, max_gap=8):
    paragraphs = []
    current = ""
    count = 0

    for i, line in enumerate(transcript):
        text = (line.get("text") or "").strip()
        if not text:
            continue
        if text.lower() in ["[applause]", "(applause)", "[music]", "(music)"]:
            continue

        current += (" " if current else "") + text
        count += len(text)

        next_line = transcript[i + 1] if i + 1 < len(transcript) else None
        long_pause = next_line and (next_line["start"] - line["start"] > max_gap)

        if count > max_chars or long_pause or not next_line:
            paragraphs.append(current.strip())
            current = ""
            count = 0

    return paragraphs


# -------- API: Get Transcript as JSON ----------
@app.get("/transcript/{video_id}")
async def get_transcript(video_id: str):
    try:
        info, transcript = extract_transcript(video_id)
        if not transcript:
            return {"status": "error", "message": "No English captions found"}

        paragraphs = merge_paragraphs(transcript)

        return {
            "status": "success",
            "video_title": info.get("title") or "Transcript",
            "language": "en",
            "transcript": [{"text": p} for p in paragraphs]
        }
    except Exception as e:
        emsg = str(e).lower()
        if "sign in" in emsg or "not a bot" in emsg or "cookies" in emsg or "private" in emsg:
            return {"status": "error", "message": "YouTube cookies may be expired/invalid. Please check /cookies/status and update cookies_*.txt on the server."}
        return {"status": "error", "message": str(e)}


# -------- TXT ----------
@app.get("/transcript/{video_id}/download/txt")
async def download_txt(video_id: str):
    try:
        info, transcript = extract_transcript(video_id)
        if not transcript:
            return Response("No English captions found", media_type="text/plain")

        paragraphs = merge_paragraphs(transcript)
        content = "\n\n".join(paragraphs)

        safe_title = "".join(c if c.isalnum() or c in " -_." else "_" for c in info.get("title", video_id))
        filename = f"{safe_title}.txt"

        return Response(
            content,
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return Response(str(e), media_type="text/plain")


# -------- PDF ----------
@app.get("/transcript/{video_id}/download/pdf")
async def download_pdf(video_id: str):
    try:
        info, transcript = extract_transcript(video_id)
        if not transcript:
            return Response("No English captions found", media_type="application/pdf")

        paragraphs = merge_paragraphs(transcript)

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(f"<b>{info.get('title', 'Transcript')}</b>", styles["Title"]))
        story.append(Spacer(1, 12))

        for p in paragraphs:
            story.append(Paragraph(p, styles["Normal"]))
            story.append(Spacer(1, 12))

        doc.build(story)
        buffer.seek(0)

        safe_title = "".join(c if c.isalnum() or c in " -_." else "_" for c in info.get("title", video_id))
        filename = f"{safe_title}.pdf"

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename=\"{filename}\"'}
        )
    except Exception as e:
        return Response(str(e), media_type="application/pdf")


# -------- DOCX ----------
@app.get("/transcript/{video_id}/download/docx")
async def download_docx(video_id: str):
    try:
        info, transcript = extract_transcript(video_id)
        if not transcript:
            return Response("No English captions found", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        paragraphs = merge_paragraphs(transcript)

        doc = Document()
        doc.add_heading(info.get("title", "Transcript"), level=1)
        for p in paragraphs:
            doc.add_paragraph(p)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        safe_title = "".join(c if c.isalnum() or c in " -_." else "_" for c in info.get("title", video_id))
        filename = f"{safe_title}.docx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename=\"{filename}\"'}
        )
    except Exception as e:
        return Response(str(e), media_type="text/plain")


# -------- COOKIE HEALTHCHECK ----------
@app.get("/cookies/status")
async def cookies_status(video_id: str = "dQw4w9WgXcQ"):
    """
    Test each cookie file by trying to fetch captions for a known public video.
    Default: Rick Astley 'Never Gonna Give You Up'.
    """
    url = f"https://www.youtube.com/watch?v={video_id}"
    results = []

    cookie_files = _find_cookie_files()

    for cookie_file in cookie_files:
        ydl_opts = {
            "skip_download": True,
            "quiet": True,
            "cookiefile": cookie_file,
        }
        try:
            info = _extract_with_options(url, ydl_opts)
            if info.get("subtitles") or info.get("automatic_captions"):
                results.append({"cookie": cookie_file, "status": "✅ working"})
            else:
                results.append({"cookie": cookie_file, "status": "⚠️ no captions found"})
        except Exception as e:
            results.append({"cookie": cookie_file, "status": f"❌ failed - {str(e)[:120]}..."})

    # Update preferred cookies (working ones go to front)
    working = [r["cookie"] for r in results if "working" in r["status"]]
    # Put working cookies first in preferred_cookies
    global preferred_cookies
    for c in reversed(working):
        if c in preferred_cookies:
            preferred_cookies.remove(c)
        preferred_cookies.insert(0, c)

    return {
        "checked_video": url,
        "results": results,
        "preferred_order": preferred_cookies
    }
